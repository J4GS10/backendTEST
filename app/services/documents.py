"""
Generación de Actas de Entrega / Hojas de Descargo en Word (.docx) y PDF.

El formato reproduce el acta corporativa de referencia (ejemplo.pdf):
  · Encabezado: logo (izq) + código de formulario y "Ciudad, fecha" (der).
  · Título centrado en negrita: "NOTA: ENTREGA…" / "NOTA: DEVOLUCIÓN A …".
  · Párrafo de respaldo + tabla DISPOSITIVO | MARCA | MODELO | SERIE con
    cabecera oscura y texto blanco.
  · Párrafos de casuística + "Descripción: <motivo>".
  · Dos firmas (Entrega / Recibe) con nombre, rol, departamento y empresa.
  · Pie de página legal en letra pequeña.

Si el logo no se puede descargar, se degrada al nombre de la empresa en texto.
"""
from __future__ import annotations

from io import BytesIO
from datetime import datetime
from typing import List, Optional
import uuid

import httpx
import structlog
from fastapi import HTTPException
from sqlalchemy.ext.asyncio import AsyncSession

from app.repositories.traceability import TraceabilityRepository
from app.repositories.governance import GovernanceRepository

log = structlog.get_logger("documents")

# Paleta corporativa del acta de referencia.
_HEADER_BG = (31, 58, 95)     # navy de la cabecera de la tabla
_INK = (15, 23, 42)
_MUTED = (90, 100, 115)
_DEFAULT_PRIMARY = (31, 58, 95)

# Datos institucionales del encabezado (configurables a futuro).
_CIUDAD = "Guatemala"
_CODIGO_FORM = "F.IT.GUA.04.01"

_MESES = ["enero", "febrero", "marzo", "abril", "mayo", "junio", "julio",
          "agosto", "septiembre", "octubre", "noviembre", "diciembre"]

_ACTA_TEXTS = {
    "entrega": {
        "titulo": "NOTA: ENTREGA DE EQUIPO",
        "intro": ("El siguiente documento se extiende como un respaldo para la entrega de "
                  "dispositivos de {empresa} al colaborador {colaborador}, en la cual se "
                  "detalla las características de los dispositivos que se entregan."),
        "cuerpo": ("Es para conocimiento y respaldo de administración, sobre la entrega de "
                   "activos de la empresa que quedan bajo la custodia y responsabilidad del "
                   "colaborador, quien se compromete a darles un uso adecuado y a devolverlos "
                   "cuando la empresa así lo requiera."),
        "firma_izq": ("Recibe equipo", "Colaborador"),
        "firma_der": ("Entrega / Autoriza", "Departamento de TI"),
    },
    "descargo": {
        "titulo": "NOTA: DEVOLUCIÓN A {empresa}",
        "intro": ("El siguiente documento se extiende como un respaldo para la devolución de "
                  "dispositivos a {empresa}, en la cual se detalla las características de los "
                  "dispositivos que se reciben."),
        "cuerpo": ("Es para conocimiento y respaldo de administración, sobre la recepción de "
                   "activos de la empresa, por casuísticas, tales como: Daño, equipo obsoleto, "
                   "equipo propio del empleado y/o caso especial autorizado."),
        "firma_izq": ("Entrega equipo", "Colaborador"),
        "firma_der": ("Recibe", "Departamento de TI"),
    },
}

_PIE_LEGAL = (
    "Este documento debe ser escaneado y enviado por correo electrónico a administración y "
    "jefaturas, relacionadas con la adquisición y retiro de equipos informáticos, como activos "
    "propios y/o ajenos de la empresa. {empresa} no es responsable del mal uso que se le pueda "
    "dar al mismo, fuera de oficinas. El documento no tiene principios legales, únicamente es "
    "para control interno y conocimiento de las partes implicadas, para uso exclusivo de {empresa}."
)


class DocumentService:
    def __init__(self, db: AsyncSession):
        self.db = db
        self.trace_repo = TraceabilityRepository(db)
        self.gov_repo = GovernanceRepository(db)

    # =================================================================
    # API PÚBLICA
    # =================================================================
    async def generar_acta_entrega(self, movimiento_id: uuid.UUID, formato: str = "docx",
                                   tipo: str = "entrega", mensajero: Optional[str] = None) -> BytesIO:
        data = await self._gather([movimiento_id], tipo=tipo, mensajero=mensajero)
        return self._render(data, formato)

    async def generar_acta_multiple(self, movimiento_ids: List[uuid.UUID], formato: str = "docx",
                                    tipo: str = "entrega", mensajero: Optional[str] = None) -> BytesIO:
        data = await self._gather(movimiento_ids, tipo=tipo, mensajero=mensajero)
        return self._render(data, formato)

    def _render(self, data: dict, formato: str) -> BytesIO:
        return self._render_pdf(data) if (formato or "").lower() == "pdf" else self._render_docx(data)

    # =================================================================
    # RECOLECCIÓN DE DATOS
    # =================================================================
    async def _gather(self, movimiento_ids: List[uuid.UUID], tipo: str = "entrega",
                      mensajero: Optional[str] = None) -> dict:
        tipo = tipo if tipo in _ACTA_TEXTS else "entrega"
        movimientos = []
        for mid in movimiento_ids:
            mov = await self.trace_repo.get_by_id_full(mid)
            if not mov:
                raise HTTPException(status_code=404, detail=f"MOVEMENT_NOT_FOUND: {mid}")
            movimientos.append(mov)
        if not movimientos:
            raise HTTPException(status_code=400, detail="NO_MOVEMENTS_PROVIDED")

        persona = movimientos[0].persona
        for mov in movimientos:
            if mov.PER_Persona != persona.PER_Persona:
                raise HTTPException(status_code=400, detail="CANNOT_MIX_DIFFERENT_PEOPLE_IN_SAME_DOCUMENT")

        config = await self.gov_repo.get_config()
        empresa = (config.SYS_Nombre_Empresa or "Mi Empresa").strip()
        logo = await self._fetch_logo(config.SYS_Logo_URL)

        items = []
        for mov in movimientos:
            a = mov.activo
            tipo_act = a.tipo_activo.TAC_Nombre if a and a.tipo_activo else ""
            marca = a.modelo.marca.MAR_Nombre if a and a.modelo and a.modelo.marca else ""
            modelo = a.modelo.MOD_Nombre if a and a.modelo else ""
            hostname = (a.ACT_Hostname if a else "") or ""
            # "DISPOSITIVO" = hostname si existe, si no el tipo, si no el código.
            dispositivo = hostname or tipo_act or (a.ACT_Codigo_Interno if a else "—")
            items.append({
                "dispositivo": dispositivo or "—",
                "marca": (marca or "—").upper(),
                "modelo": modelo or "-",
                "serie": (a.ACT_Serie_Fabricante if a else "") or "-",
            })

        if tipo == "descargo" and movimientos[0].MOV_Fecha_Devolucion:
            d = movimientos[0].MOV_Fecha_Devolucion
        else:
            d = datetime.now()
        motivo_default = "Devolución de equipo." if tipo == "descargo" else "Asignación de equipo."

        return {
            "tipo": tipo,
            "txt": _ACTA_TEXTS[tipo],
            "empresa": empresa,
            "logo": logo,
            "primary": self._hex_to_rgb(config.SYS_Color_Primario, _DEFAULT_PRIMARY),
            "colaborador": f"{persona.PER_Primer_Nombre} {persona.PER_Primer_Apellido}".strip(),
            "mensajero": (mensajero or "").strip() or None,
            "ciudad": (config.SYS_Ciudad or _CIUDAD).strip(),
            "codigo_form": (config.SYS_Codigo_Formulario or _CODIGO_FORM).strip(),
            "fecha_larga": f"{d.day} de {_MESES[d.month - 1]} de {d.year}",
            "acta_ref": str(movimientos[0].MOV_Movimiento)[:8].upper(),
            "motivo": movimientos[0].MOV_Observacion or motivo_default,
            "items": items,
        }

    async def _fetch_logo(self, url: Optional[str]) -> Optional[bytes]:
        if not url or not isinstance(url, str) or not url.lower().startswith(("http://", "https://")):
            return None
        try:
            async with httpx.AsyncClient(timeout=4.0, follow_redirects=True) as client:
                r = await client.get(url)
            ct = r.headers.get("content-type", "")
            if r.status_code == 200 and ct.startswith("image/") and 0 < len(r.content) <= 3_000_000:
                return r.content
        except Exception as e:  # noqa: BLE001
            log.warning("documents.logo_fetch_failed", error=str(e)[:160])
        return None

    @staticmethod
    def _hex_to_rgb(h: Optional[str], default: tuple = _DEFAULT_PRIMARY) -> tuple:
        s = (h or "").lstrip("#")
        try:
            if len(s) == 6:
                return tuple(int(s[i:i + 2], 16) for i in (0, 2, 4))
            if len(s) == 3:
                return tuple(int(s[i] * 2, 16) for i in range(3))
        except ValueError:
            pass
        return default

    @staticmethod
    def _logo_size_inches(img_bytes: bytes, max_w: float, max_h: float) -> Optional[tuple]:
        try:
            from PIL import Image as PILImage
            with PILImage.open(BytesIO(img_bytes)) as im:
                w, h = im.size
            if w <= 0 or h <= 0:
                return None
            scale = min(max_w / w, max_h / h)
            return (w * scale, h * scale)
        except Exception:  # noqa: BLE001
            return None

    # =================================================================
    # RENDER WORD (.docx)
    # =================================================================
    def _render_docx(self, data: dict) -> BytesIO:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_ALIGN_VERTICAL
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        empresa = data["empresa"]
        txt = data["txt"]
        header_hex = "%02X%02X%02X" % _HEADER_BG

        def shade(cell, hex_color):
            tcPr = cell._tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), hex_color)
            tcPr.append(shd)

        def no_borders(tbl):
            borders = OxmlElement("w:tblBorders")
            for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
                e = OxmlElement(f"w:{edge}"); e.set(qn("w:val"), "none"); borders.append(e)
            tbl._tbl.tblPr.append(borders)

        doc = Document()
        st = doc.styles["Normal"]; st.font.name = "Arial"; st.font.size = Pt(11)
        for s in doc.sections:
            s.top_margin = Inches(0.8); s.bottom_margin = Inches(0.7)
            s.left_margin = Inches(0.9); s.right_margin = Inches(0.9)

        # --- ENCABEZADO: logo (izq) + código/fecha (der) ---
        head = doc.add_table(rows=1, cols=2); no_borders(head)
        head.columns[0].width = Inches(3.4); head.columns[1].width = Inches(3.4)
        cl, cr = head.rows[0].cells
        cl.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cr.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        pl = cl.paragraphs[0]; pl.alignment = WD_ALIGN_PARAGRAPH.LEFT
        placed = False
        if data["logo"]:
            size = self._logo_size_inches(data["logo"], max_w=2.6, max_h=1.0)
            if size:
                try:
                    pl.add_run().add_picture(BytesIO(data["logo"]), width=Inches(size[0]), height=Inches(size[1]))
                    placed = True
                except Exception:  # noqa: BLE001
                    placed = False
        if not placed:
            r = pl.add_run(empresa); r.bold = True; r.font.size = Pt(22); r.font.color.rgb = RGBColor(*_HEADER_BG)
        pr = cr.paragraphs[0]; pr.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r1 = pr.add_run(data["codigo_form"]); r1.bold = True; r1.font.size = Pt(10)
        r2 = pr.add_run(f"\n{data['ciudad']} {data['fecha_larga']}"); r2.font.size = Pt(10)

        doc.add_paragraph()
        # --- TÍTULO ---
        ti = doc.add_paragraph(); ti.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rt = ti.add_run(txt["titulo"].format(empresa=empresa.upper())); rt.bold = True; rt.font.size = Pt(12)
        doc.add_paragraph()

        # --- INTRO ---
        intro = doc.add_paragraph(); intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        intro.add_run(txt["intro"].format(empresa=f"{empresa} S.A.", colaborador=data["colaborador"]))
        doc.add_paragraph()

        # --- TABLA ---
        headers = ["DISPOSITIVO", "MARCA", "MODELO", "SERIE"]
        table = doc.add_table(rows=1, cols=4); table.style = "Table Grid"
        table.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for i, h in enumerate(headers):
            c = table.rows[0].cells[i]; shade(c, header_hex)
            p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            rr = p.add_run(h); rr.bold = True; rr.font.size = Pt(10); rr.font.color.rgb = RGBColor(255, 255, 255)
        for it in data["items"]:
            cells = table.add_row().cells
            for i, key in enumerate(("dispositivo", "marca", "modelo", "serie")):
                p = cells[i].paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.add_run(str(it[key])).font.size = Pt(10)
        doc.add_paragraph()

        # --- CUERPO ---
        cu = doc.add_paragraph(); cu.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        cu.add_run(txt["cuerpo"])
        de = doc.add_paragraph(); de.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        de.add_run("Descripción: ").bold = True
        de.add_run(data["motivo"])

        for _ in range(5):
            doc.add_paragraph()

        # --- FIRMAS ---
        fz, fd = txt["firma_izq"], txt["firma_der"]
        # Firma derecha: Departamento de TI; si hay mensajero externo, lo recibe él.
        if data.get("mensajero"):
            der_sig = (data["mensajero"], fd[0], "Mensajero externo", f"No pertenece a {empresa}")
        else:
            der_sig = ("", fd[0], fd[1], empresa)
        ft = doc.add_table(rows=1, cols=2); no_borders(ft)
        izq, der = ft.rows[0].cells
        for cell, (nombre, accion, depto, comp) in (
            (izq, (data["colaborador"], fz[0], fz[1], empresa)),
            (der, der_sig),
        ):
            p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run("_______________________________\n").bold = False
            if nombre:
                rn = p.add_run(nombre + "\n"); rn.bold = True; rn.font.size = Pt(10.5)
            p.add_run(accion + "\n").font.size = Pt(10)
            rd = p.add_run(depto + "\n"); rd.bold = True; rd.font.size = Pt(10)
            re = p.add_run(comp); re.font.size = Pt(9.5); re.font.color.rgb = RGBColor(*_MUTED)

        for _ in range(2):
            doc.add_paragraph()
        # --- PIE LEGAL ---
        pie = doc.add_paragraph(); pie.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        rp = pie.add_run("•  " + _PIE_LEGAL.format(empresa=empresa))
        rp.font.size = Pt(7.5); rp.font.color.rgb = RGBColor(*_MUTED)

        buf = BytesIO(); doc.save(buf); buf.seek(0)
        return buf

    # =================================================================
    # RENDER PDF (reportlab)
    # =================================================================
    def _render_pdf(self, data: dict) -> BytesIO:
        from reportlab.lib.pagesizes import LETTER
        from reportlab.lib.units import inch
        from reportlab.lib import colors
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.enums import TA_JUSTIFY, TA_RIGHT, TA_CENTER, TA_LEFT
        from reportlab.platypus import (
            SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image as RLImage,
        )

        empresa = data["empresa"]
        txt = data["txt"]
        header_bg = colors.Color(*[c / 255 for c in _HEADER_BG])
        ink = colors.Color(*[c / 255 for c in _INK])
        muted = colors.Color(*[c / 255 for c in _MUTED])

        ss = getSampleStyleSheet()
        st_code = ParagraphStyle("code", parent=ss["Normal"], fontName="Helvetica-Bold", fontSize=10,
                                 alignment=TA_RIGHT, textColor=ink, leading=14)
        st_date = ParagraphStyle("date", parent=ss["Normal"], fontSize=10, alignment=TA_RIGHT, textColor=ink)
        st_title = ParagraphStyle("title", parent=ss["Normal"], fontName="Helvetica-Bold", fontSize=12,
                                  alignment=TA_CENTER, textColor=ink, leading=16)
        st_body = ParagraphStyle("body", parent=ss["Normal"], fontName="Helvetica", fontSize=10.5,
                                 alignment=TA_JUSTIFY, leading=15)
        st_cellh = ParagraphStyle("ch", parent=ss["Normal"], fontName="Helvetica-Bold", fontSize=9.5,
                                  textColor=colors.white, alignment=TA_CENTER, leading=12)
        st_cell = ParagraphStyle("cl", parent=ss["Normal"], fontSize=9.5, alignment=TA_CENTER, leading=12)
        st_sign = ParagraphStyle("sg", parent=ss["Normal"], fontSize=10, alignment=TA_CENTER, leading=14)
        st_foot = ParagraphStyle("ft", parent=ss["Normal"], fontSize=7.5, alignment=TA_JUSTIFY,
                                 textColor=muted, leading=10)
        st_brand = ParagraphStyle("br", parent=ss["Normal"], fontName="Helvetica-Bold", fontSize=22,
                                  textColor=header_bg, alignment=TA_LEFT)

        story = []

        # --- ENCABEZADO ---
        logo_flow = None
        if data["logo"]:
            size = self._logo_size_inches(data["logo"], max_w=2.6, max_h=1.0)
            if size:
                try:
                    logo_flow = RLImage(BytesIO(data["logo"]), width=size[0] * inch, height=size[1] * inch)
                    logo_flow.hAlign = "LEFT"
                except Exception:  # noqa: BLE001
                    logo_flow = None
        if logo_flow is None:
            logo_flow = Paragraph(empresa, st_brand)
        right = [Paragraph(data["codigo_form"], st_code),
                 Paragraph(f"{data['ciudad']} {data['fecha_larga']}", st_date)]
        header = Table([[logo_flow, right]], colWidths=[3.4 * inch, 3.4 * inch])
        header.setStyle(TableStyle([("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                                    ("LEFTPADDING", (0, 0), (-1, -1), 0),
                                    ("RIGHTPADDING", (0, 0), (-1, -1), 0)]))
        story.append(header)
        story.append(Spacer(1, 22))

        # --- TÍTULO ---
        story.append(Paragraph(txt["titulo"].format(empresa=empresa.upper()), st_title))
        story.append(Spacer(1, 18))

        # --- INTRO ---
        story.append(Paragraph(
            txt["intro"].format(empresa=f"{empresa} S.A.", colaborador=f"<b>{data['colaborador']}</b>"),
            st_body))
        story.append(Spacer(1, 14))

        # --- TABLA ---
        tdata = [[Paragraph(h, st_cellh) for h in ("DISPOSITIVO", "MARCA", "MODELO", "SERIE")]]
        for it in data["items"]:
            tdata.append([Paragraph(str(it["dispositivo"]), st_cell),
                          Paragraph(str(it["marca"]), st_cell),
                          Paragraph(str(it["modelo"]), st_cell),
                          Paragraph(str(it["serie"]), st_cell)])
        tbl = Table(tdata, colWidths=[1.9 * inch, 1.35 * inch, 1.85 * inch, 1.65 * inch], repeatRows=1)
        tbl.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), header_bg),
            ("GRID", (0, 0), (-1, -1), 0.7, colors.Color(0.55, 0.6, 0.68)),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ("TOPPADDING", (0, 0), (-1, -1), 7), ("BOTTOMPADDING", (0, 0), (-1, -1), 7),
        ]))
        story.append(tbl)
        story.append(Spacer(1, 16))

        # --- CUERPO ---
        story.append(Paragraph(txt["cuerpo"], st_body))
        story.append(Spacer(1, 10))
        story.append(Paragraph(f"<b>Descripción:</b> {data['motivo']}", st_body))
        story.append(Spacer(1, 66))

        # --- FIRMAS ---
        fz, fd = txt["firma_izq"], txt["firma_der"]
        co_style = ParagraphStyle("co", parent=st_sign, fontSize=9.5, textColor=muted)

        def firma(nombre, accion, depto, comp):
            out = [Paragraph("_______________________________", st_sign)]
            if nombre:
                out.append(Paragraph(f"<b>{nombre}</b>", st_sign))
            out.append(Paragraph(accion, st_sign))
            out.append(Paragraph(f"<b>{depto}</b>", st_sign))
            out.append(Paragraph(comp, co_style))
            return out

        if data.get("mensajero"):
            der_sig = firma(data["mensajero"], fd[0], "Mensajero externo", f"No pertenece a {empresa}")
        else:
            der_sig = firma("", fd[0], fd[1], empresa)
        firmas = Table([[firma(data["colaborador"], fz[0], fz[1], empresa), der_sig]],
                       colWidths=[3.35 * inch, 3.35 * inch])
        firmas.setStyle(TableStyle([("VALIGN", (0, 0), (-1, -1), "TOP")]))
        story.append(firmas)
        story.append(Spacer(1, 30))

        # --- PIE LEGAL ---
        story.append(Paragraph("•&nbsp;&nbsp;" + _PIE_LEGAL.format(empresa=empresa), st_foot))

        buf = BytesIO()
        pdf = SimpleDocTemplate(buf, pagesize=LETTER, topMargin=0.8 * inch, bottomMargin=0.7 * inch,
                                leftMargin=0.9 * inch, rightMargin=0.9 * inch,
                                title=f"{txt['titulo'].format(empresa=empresa)} {data['acta_ref']}",
                                author=empresa)
        pdf.build(story)
        buf.seek(0)
        return buf
