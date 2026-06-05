"""Tests de actas (entrega/descargo) en Word y PDF, formato corporativo."""
from io import BytesIO

from docx import Document

from app.services.documents import DocumentService, _ACTA_TEXTS


def _data(logo=None, tipo="entrega"):
    return {
        "tipo": tipo, "txt": _ACTA_TEXTS[tipo],
        "empresa": "Lombardi", "logo": logo, "primary": (31, 58, 95),
        "colaborador": "Jose Gonzalez", "ciudad": "Guatemala",
        "codigo_form": "F.IT.GUA.04.01", "fecha_larga": "11 de marzo de 2026",
        "acta_ref": "AB12CD34", "motivo": "Finalidad de labores. Uso normal.",
        "items": [
            {"dispositivo": "DESKTOP-KJEXT4", "marca": "DELL", "modelo": "XPS-8920", "serie": "7XD0JH2"},
            {"dispositivo": "TECLADO/MOUSE", "marca": "MICROSOFT", "modelo": "3J2-0008", "serie": "-"},
        ],
    }


def _png_logo() -> bytes:
    from PIL import Image
    im = Image.new("RGB", (200, 60), (31, 58, 95))
    b = BytesIO(); im.save(b, "PNG"); return b.getvalue()


def _svc() -> DocumentService:
    return DocumentService.__new__(DocumentService)


def _docx_text(buf: BytesIO) -> str:
    doc = Document(buf)
    return "\n".join(p.text for p in doc.paragraphs) + "\n" + "\n".join(
        c.text for tbl in doc.tables for row in tbl.rows for c in row.cells)


def test_acta_entrega_docx_valida():
    out = _svc()._render_docx(_data(tipo="entrega")).getvalue()
    assert out[:2] == b"PK" and len(out) > 5000


def test_acta_entrega_pdf_valida():
    out = _svc()._render_pdf(_data(tipo="entrega")).getvalue()
    assert out[:4] == b"%PDF" and len(out) > 1500


def test_formato_corporativo_docx():
    """El docx reproduce el formato de referencia: cabecera DISPOSITIVO, firmas, pie."""
    texto = _docx_text(_svc()._render_docx(_data(tipo="descargo")))
    assert "DEVOLUCIÓN A LOMBARDI" in texto       # título
    assert "DISPOSITIVO" in texto and "SERIE" in texto  # cabecera de tabla
    assert "DESKTOP-KJEXT4" in texto               # columna dispositivo (hostname)
    assert "Entrega equipo" in texto and "Recibe" in texto  # firmas
    assert "Descripción:" in texto
    assert "control interno" in texto              # pie legal


def test_hoja_descargo_pdf_valida():
    out = _svc()._render_pdf(_data(tipo="descargo")).getvalue()
    assert out[:4] == b"%PDF" and len(out) > 1500


def test_acta_con_mensajero_externo():
    """Si se indica un mensajero, recibe él (externo, no pertenece a la empresa)."""
    d = _data(tipo="descargo")
    d["mensajero"] = "Pedro Mensajero"
    texto = _docx_text(_svc()._render_docx(d))
    assert "Pedro Mensajero" in texto
    assert "Mensajero externo" in texto
    assert "No pertenece a Lombardi" in texto
    assert _svc()._render_pdf(d).getvalue()[:4] == b"%PDF"


def test_acta_con_logo_no_falla():
    logo = _png_logo()
    assert _svc()._render_docx(_data(logo, "entrega")).getvalue()[:2] == b"PK"
    assert _svc()._render_pdf(_data(logo, "descargo")).getvalue()[:4] == b"%PDF"


def test_acta_sin_logo_degrada():
    assert _svc()._render_docx(_data(None)).getvalue()[:2] == b"PK"
    assert _svc()._render_pdf(_data(None)).getvalue()[:4] == b"%PDF"


def test_hex_to_rgb():
    assert DocumentService._hex_to_rgb("#1f3a5f") == (31, 58, 95)
    assert DocumentService._hex_to_rgb("#abc") == (170, 187, 204)
    assert DocumentService._hex_to_rgb("invalido") == (31, 58, 95)
